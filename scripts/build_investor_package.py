from pathlib import Path
import json
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("investor-package")
DATA_ROOM = ROOT / "data-room"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
INK = "111827"
MUTED = "4B5563"
AMBER = "7A5A00"

METRICS = {
    "as_of": "August 6, 2026",
    "users": 120,
    "events": 45,
    "registrations": 482,
    "active_events": 13,
    "organizers_with_events": 27,
    "organizers_with_30_plus_confirmed": 4,
    "organizers_last_30_days": 5,
    "pro_users": 1,
    "free_users": 119,
    "report_revenue_ksh": 0,
    "report_downloads": 0,
    "top_events": [
        {"title": "DISRUPTORS CONVECTION 7", "confirmed_count": 92},
        {
            "title": "IspeakDisruptors Convention 6 Disruptors Convention | Systemic Influence",
            "confirmed_count": 91,
        },
        {"title": "Regina Daniels Children's Day Carnival (Spelling Bee)", "confirmed_count": 50},
        {"title": "Tum Sports Tournament", "confirmed_count": 38},
        {"title": "Event 002 (Copy)", "confirmed_count": 37},
    ],
}

ROUND = {
    "raise_usd": 100000,
    "equity_pct": 10,
    "pre_money_usd": 900000,
    "post_money_usd": 1000000,
    "runway_claim": "12 months base, stretchable toward 18 months with disciplined execution",
}

USE_OF_FUNDS = [
    ("Product and engineering", "USD 28,000", "Strengthen core product, reliability, bug fixing, and payment/event operations features."),
    ("Growth and customer acquisition", "USD 25,000", "Founder-led sales, field marketing, partnerships, demos, onboarding, and campaigns."),
    ("Customer support and operations", "USD 14,000", "Support capacity, helpdesk processes, onboarding material, event-day issue handling."),
    ("Security and compliance", "USD 12,000", "Security retainer, penetration testing, privacy/data-protection work, and governance basics."),
    ("Cloud and infrastructure", "USD 13,000", "Hosting, database, storage, email delivery, monitoring, backups, and scaling headroom."),
    ("Reserve and working capital", "USD 8,000", "Runway protection, finance admin, legal, and unforeseen operating pressure."),
]


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_run(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D1D5DB", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def base_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.167
    return doc


def header_footer(doc: Document, header_text: str):
    header = doc.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run(header_text)
    set_run(r, size=9, color=MUTED)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Prepared August 2026")
    set_run(r, size=9, color=MUTED)


def para(doc, text="", size=11, bold=False, color=INK, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, DARK_BLUE, 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run(r, size=size, bold=True, color=color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.167
    r = p.add_run(text)
    set_run(r, size=10.5, color=INK)


def callout(doc, title, body, fill=CALLOUT):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t, color="BFD3EA")
    cell = t.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=160, bottom=160, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_table_borders(t)
    for row in t.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=font_size, bold=True, color=INK)
    for row_values in rows:
        row = t.add_row().cells
        for i, value in enumerate(row_values):
            set_cell_margins(row[i])
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = row[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=font_size, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def write_summary():
    doc = base_doc()
    header_footer(doc, "EventSlot Executive Summary")
    para(doc, "EXECUTIVE SUMMARY", size=10, bold=True, color=BLUE, after=3)
    para(doc, "EventSlot", size=24, bold=True, color="000000", after=4)
    para(doc, "A working event operations platform for organizers.", size=12.5, color=MUTED, after=12)
    callout(
        doc,
        "Round summary",
        "EventSlot is raising USD 100,000 at a USD 900,000 pre-money valuation in exchange for 10% equity. "
        "The round is designed to fund 12 months of focused execution, with the potential to stretch toward 18 months through disciplined spend.",
        fill=LIGHT_BLUE,
    )
    para(
        doc,
        "EventSlot helps organizers manage registration, capacity, attendee flow, communication, check-in, and event delivery from one place instead of juggling separate tools such as WhatsApp, forms, spreadsheets, manual payment tracking, and paper attendance lists.",
    )
    para(
        doc,
        f"As of {METRICS['as_of']}, EventSlot has {METRICS['users']} users, {METRICS['events']} events created, {METRICS['registrations']} registrations processed, {METRICS['active_events']} active events, and {METRICS['organizers_with_events']} organizers who have created events. "
        f"{METRICS['organizers_with_30_plus_confirmed']} organizers have already run events with 30 or more confirmed attendees.",
    )
    para(
        doc,
        "The opportunity is not being presented as a mature revenue business. It is being presented as a working product with real usage, now raising capital to convert early adoption into repeatable growth, stronger organizer retention, and paid conversion.",
    )
    heading(doc, "What This Round Funds")
    for name, amount, description in USE_OF_FUNDS:
        bullet(doc, f"{name}: {amount}. {description}")
    heading(doc, "12-Month Objectives")
    bullet(doc, "Increase active organizers and repeat event usage.")
    bullet(doc, "Convert early organizer demand into paid plans and premium feature revenue.")
    bullet(doc, "Deepen product reliability, support quality, and security discipline.")
    bullet(doc, "Build a more repeatable go-to-market motion in focused customer segments.")
    bullet(doc, "Prepare EventSlot for a stronger follow-on raise based on clearer commercial traction.")
    doc.save(ROOT / "EventSlot_Executive_Summary.docx")


def write_investment_memo():
    doc = base_doc()
    header_footer(doc, "EventSlot Investment Memo")
    para(doc, "DETAILED INVESTMENT MEMO", size=10, bold=True, color=BLUE, after=3)
    para(doc, "EventSlot Investment Opportunity", size=24, bold=True, color="000000", after=4)
    para(doc, "Working product, live organizer usage, and a focused 12-month execution plan.", size=12.5, color=MUTED, after=12)
    callout(
        doc,
        "Investment terms",
        "Raise: USD 100,000 | Pre-money valuation: USD 900,000 | Post-money valuation: USD 1,000,000 | Equity: 10%",
        fill=LIGHT_BLUE,
    )
    heading(doc, "1. Company Overview")
    para(
        doc,
        "EventSlot is a digital event operations platform that helps organizers create events, accept registrations, manage capacity, coordinate attendees, support verification and check-in, and review event performance. The company is based in Nairobi, Kenya and is currently at pre-seed stage.",
    )
    heading(doc, "2. Problem")
    para(
        doc,
        "Many organizers still run one event across several disconnected systems. Registration may happen in Google Forms, communication in WhatsApp, payment confirmation in mobile money statements, attendee updates in spreadsheets, and check-in on paper or manual lists. This fragmentation creates confusion, slow response times, and weak operational visibility.",
    )
    bullet(doc, "Organizers struggle to know who registered, who paid, who is confirmed, and how much capacity remains.")
    bullet(doc, "Manual admin work consumes time that should go into growth, guest experience, and event quality.")
    bullet(doc, "Attendees experience delayed confirmations, unclear status, and inefficient event-day flow.")
    heading(doc, "3. Product")
    para(
        doc,
        "EventSlot replaces fragmented event administration with one connected workflow. Organizers create the event, share one registration link, manage attendees from one dashboard, and deliver the event with stronger control over capacity, communication, and attendance operations.",
    )
    table(
        doc,
        ["Product area", "What it does"],
        [
            ["Event creation", "Create events, registration questions, deadlines, locations, and public pages."],
            ["Registration flow", "Collect attendee details through a dedicated event link and structured form."],
            ["Capacity and waitlist", "Control available slots and manage attendee flow when events fill up."],
            ["Verification and check-in", "Support attendee lookup, ticket verification, and event-day attendance handling."],
            ["Analytics and exports", "Track registrations and event activity with reporting surfaces and exports."],
            ["Organizer tools", "Support communication, team access, and event operations from one place."],
        ],
        widths=[1.9, 4.6],
    )
    heading(doc, "4. Current Traction")
    para(doc, f"Verified live platform snapshot as of {METRICS['as_of']}:")
    table(
        doc,
        ["Metric", "Value", "Meaning"],
        [
            ["Users", METRICS["users"], "Real product adoption, though still early."],
            ["Events created", METRICS["events"], "Evidence of actual organizer use."],
            ["Registrations processed", METRICS["registrations"], "Proof that real attendee workflows are happening."],
            ["Active events", METRICS["active_events"], "Current live platform activity."],
            ["Organizers with events", METRICS["organizers_with_events"], "Early core customer base."],
            ["Organizers with 30+ confirmed attendees on an event", METRICS["organizers_with_30_plus_confirmed"], "Initial signal of meaningful event usage."],
            ["Paid users", METRICS["pro_users"], "Monetization exists but is still very early."],
        ],
        widths=[2.35, 0.95, 3.2],
    )
    heading(doc, "5. Customer Focus")
    para(
        doc,
        "EventSlot can eventually serve many event categories, but the early go-to-market should stay focused on organizer groups with repeated coordination pain and a clear need for structured registration and attendance management.",
    )
    bullet(doc, "Churches and ministry teams")
    bullet(doc, "Universities and student organizations")
    bullet(doc, "Training providers and workshop organizers")
    bullet(doc, "Communities and independent event organizers")
    heading(doc, "6. Business Model")
    para(
        doc,
        "EventSlot is expected to monetize through a mix of paid organizer plans, premium features, and transaction-linked revenue for paid event workflows. The current raise is intended to validate which revenue streams convert best and how quickly organizer usage turns into commercial value.",
    )
    bullet(doc, "Paid organizer subscriptions")
    bullet(doc, "Premium workflow features")
    bullet(doc, "Paid-event transaction-linked revenue")
    bullet(doc, "Institutional or team packages over time")
    heading(doc, "7. Why This Round Matters")
    para(
        doc,
        "This round is not about funding a fully mature company. It is about helping EventSlot cross the gap from early usage to repeatable commercial traction. The product already works. The next step is to prove that organizers come back, pay, and grow with the platform.",
    )
    heading(doc, "8. Use of Funds")
    table(
        doc,
        ["Area", "Allocation", "Purpose"],
        USE_OF_FUNDS,
        widths=[2.05, 1.15, 3.3],
    )
    heading(doc, "9. Milestones For The Next 12 Months")
    bullet(doc, "Increase the number of active organizers and repeat-event creators.")
    bullet(doc, "Convert early demand into more paid plans and revenue visibility.")
    bullet(doc, "Improve onboarding, support, and response speed for organizers and attendees.")
    bullet(doc, "Strengthen security, compliance, and operational reliability.")
    bullet(doc, "Build a stronger fundraising case for a follow-on round based on validated growth.")
    heading(doc, "10. Risks And Mitigation")
    bullet(doc, "Low paid conversion risk: keep pricing simple and closely observe who asks for premium workflows.")
    bullet(doc, "Slow organizer acquisition risk: focus on beachhead segments and founder-led selling before broad expansion.")
    bullet(doc, "Operational risk on live events: invest early in support, verification reliability, and security discipline.")
    bullet(doc, "Infrastructure and payment complexity risk: keep cloud, backups, logs, and payment operations visible from day one.")
    doc.save(ROOT / "EventSlot_Detailed_Investment_Memo.docx")


def write_data_room_doc():
    doc = base_doc()
    header_footer(doc, "EventSlot Data Room Checklist")
    para(doc, "DATA ROOM CHECKLIST", size=10, bold=True, color=BLUE, after=3)
    para(doc, "EventSlot Investor Data Room", size=24, bold=True, color="000000", after=4)
    para(doc, "Checklist and index for the supporting materials behind the investor deck.", size=12.5, color=MUTED, after=12)
    callout(
        doc,
        "Purpose",
        "The data room should allow an investor to verify the main claims in the deck quickly: what EventSlot does, what traction exists, how the round is structured, and what the capital will fund.",
    )
    heading(doc, "1. Core Documents")
    table(
        doc,
        ["Document", "Status", "Notes"],
        [
            ["Executive summary", "Prepared", "Forwardable one-page overview."],
            ["Pitch deck", "Prepared", "Primary investor presentation."],
            ["Detailed investment memo", "Prepared", "Longer written explanation of the opportunity."],
            ["Valuation memo", "Available", "Earlier working note; use carefully so round terms stay aligned."],
            ["3-year operating budget", "Available", "Long-term planning document; not the same as this round's 12-month use-of-funds story."],
        ],
        widths=[2.1, 0.95, 3.45],
    )
    heading(doc, "2. Evidence We Should Include")
    bullet(doc, "Traction snapshot export with users, events, registrations, active events, organizer counts, and top events.")
    bullet(doc, "Product screenshots or short walkthrough captures showing event creation, registration, attendee management, verification, and analytics.")
    bullet(doc, "Use-of-funds model for the USD 100,000 round.")
    bullet(doc, "Any customer interviews, testimonials, pilot letters, or usage stories.")
    bullet(doc, "Legal/company basics such as registration docs, founder bios, and a simple cap table.")
    heading(doc, "3. Suggested Folder Structure")
    table(
        doc,
        ["Folder", "Contents"],
        [
            ["01-company", "Company registration, founder profiles, cap table draft."],
            ["02-product", "Screenshots, product summary, roadmap."],
            ["03-traction", "Live metrics export, top-event examples, monthly snapshots if available."],
            ["04-financials", "Use of funds, budget notes, valuation note, revenue model assumptions."],
            ["05-legal-and-compliance", "Terms, privacy, data-handling material, payment/compliance notes."],
            ["06-fundraise", "Executive summary, deck, memo, and meeting notes."],
        ],
        widths=[1.85, 4.7],
    )
    doc.save(ROOT / "EventSlot_Data_Room_Checklist.docx")


def write_text_artifacts():
    ensure_dir(ROOT)
    ensure_dir(DATA_ROOM)
    with open(DATA_ROOM / "traction_snapshot.json", "w", encoding="utf-8") as f:
        json.dump({"metrics": METRICS, "round": ROUND, "use_of_funds": USE_OF_FUNDS}, f, indent=2)
    with open(DATA_ROOM / "README.txt", "w", encoding="utf-8") as f:
        f.write(
            "EventSlot investor data room\n"
            "1. Use traction_snapshot.json as the live metrics reference for the current package.\n"
            "2. Add product screenshots to 02-product.\n"
            "3. Add company registration, founder bios, and cap table material to 01-company.\n"
            "4. Add any customer interviews, pilots, and proof points before sharing externally.\n"
        )


def main():
    ensure_dir(ROOT)
    write_summary()
    write_investment_memo()
    write_data_room_doc()
    write_text_artifacts()


if __name__ == "__main__":
    main()
