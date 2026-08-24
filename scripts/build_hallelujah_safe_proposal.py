from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "EventSlot_x_Hallelujah_Challenge_Proposal_safe_claims.docx"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(
    doc: Document,
    text: str = "",
    *,
    size=11,
    color="000000",
    bold=False,
    after=8,
    before=0,
    align=WD_ALIGN_PARAGRAPH.LEFT,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color="2E74B5", bold=True)
    elif level == 2:
        set_run_font(run, size=13, color="2E74B5", bold=True)
    else:
        set_run_font(run, size=12, color="1F4D78", bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        set_run_font(run, size=11, color="000000")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        set_run_font(run, size=11, color="000000")


def format_table(table, widths: list[float], header_fill="F2F4F7") -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    set_run_font(run, size=10.5, color="000000")
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("EVENTSLOT X HALLELUJAH CHALLENGE")
    set_run_font(run, size=11, color="666666", bold=True)

    add_para(doc, "Safe-claims partnership proposal", size=24, color="000000", after=2)
    add_para(
        doc,
        "This version only states features that are live today, clearly partial today, or explicitly proposed for future work.",
        size=11,
        color="555555",
        after=14,
    )

    intro = doc.add_table(rows=1, cols=1)
    intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    intro.autofit = False
    cell = intro.rows[0].cells[0]
    set_cell_width(cell, 6.5)
    shade_cell(cell, "F4F6F9")
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.12
    run = para.add_run(
        "EventSlot is an event operations platform that already handles public event pages, capacity limits, waitlists, attendance verification, team access, analytics, and email-first communication. For Hallelujah Challenge, the honest proposal is not that EventSlot already does everything, but that it can support the core registration and gate-management workflow now, while any broader workflow gets scoped as a pilot or future enhancement."
    )
    set_run_font(run, size=11.2, color="1F3A5F", bold=True)

    add_heading(doc, "1. What Hallelujah Challenge Needs", 1)
    add_para(
        doc,
        "The event pattern appears to need controlled registration, separate day-level or event-level capacity, gate verification, volunteer or team roles, and clear attendance reporting. Those are the parts EventSlot is closest to supporting today.",
    )

    add_heading(doc, "2. What EventSlot Already Offers", 1)
    add_bullets(
        doc,
        [
            "Public event pages with custom registration questions.",
            "Capacity management with confirmed-count tracking.",
            "Automatic waitlist handling when capacity is reached.",
            "FIFO promotion when capacity is increased.",
            "Ticket verification, QR flows, and walk-in check-in support.",
            "Team access for organisers and operational staff.",
            "Event analytics, exports, and recent activity visibility.",
            "Installable mobile web app behavior through PWA support.",
            "Google Calendar support and Google Maps direction support.",
            "Consent and privacy controls aligned to data-protection handling.",
            "Email-first attendee and organiser communication workflows.",
        ],
    )

    add_heading(doc, "3. Fit Check: What We Can Say Safely", 1)
    fit = doc.add_table(rows=1, cols=4)
    fit.rows[0].cells[0].text = "Event need"
    fit.rows[0].cells[1].text = "Can we claim it today?"
    fit.rows[0].cells[2].text = "Safe wording"
    fit.rows[0].cells[3].text = "Status"
    rows = [
        ("Public registration pages", "Yes", "EventSlot creates public event pages with registration", "Live"),
        ("Per-day capacity control", "Yes, by configuration", "Separate event pages or pools can control each day", "Live / configurable"),
        ("Automatic waitlist", "Yes", "Overflow registrations move to waitlist", "Live"),
        ("Gate check-in and verification", "Yes", "QR/ticket verification and walk-in check-in are available", "Live"),
        ("Volunteer or team roles", "Partially", "Team access exists; role design may need tailoring", "Live / partial"),
        ("Email notifications", "Yes", "Email-first confirmations and updates are live", "Live"),
        ("SMS or WhatsApp notifications", "No", "These are roadmap channels, not core today", "Not live"),
        ("Live broadcast", "No", "This should be presented as a proposed pilot only", "Not live"),
        ("Full multi-day event orchestration", "Partially", "Supported through event structure and configuration, but not as a special native workflow", "Partial"),
    ]
    for row_data in rows:
        row = fit.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    format_table(fit, [1.65, 1.15, 2.25, 1.45])

    add_heading(doc, "4. What We Should Not Claim", 1)
    add_bullets(
        doc,
        [
            "Do not claim live broadcast is already built.",
            "Do not claim SMS or WhatsApp notifications are native live channels.",
            "Do not imply EventSlot already has a special Hallelujah Challenge mode unless we build one.",
            "Do not claim volunteer management beyond the team-access and registration tools we already have.",
            "Do not present speculative ideas as shipped product.",
        ],
    )

    add_heading(doc, "5. Recommended Partnership Positioning", 1)
    add_para(
        doc,
        "The safest and strongest position is to say that EventSlot can support the registration and attendance backbone of Hallelujah Challenge now, with any broader workflows proposed as a pilot. That keeps the proposal useful without overselling the product.",
    )
    add_bullets(
        doc,
        [
            "Use EventSlot for public registration pages and capacity control.",
            "Use waitlist automation where registration demand exceeds capacity.",
            "Use QR or ticket verification for event-day entry.",
            "Use team access for internal operations staff.",
            "Use email communication for confirmations and updates.",
            "Treat broadcast, SMS, and WhatsApp automation as future work or separate scope.",
        ],
    )

    add_heading(doc, "6. Practical Pilot Scope", 1)
    add_numbers(
        doc,
        [
            "Start with one event day or one event series.",
            "Map the actual registration flow used by the Hallelujah Challenge team.",
            "Configure capacity, waitlist, and verification on EventSlot.",
            "Train the gate team and the core organiser team.",
            "Review attendance data after the event and decide whether to expand scope.",
        ],
    )

    add_heading(doc, "7. Honest Closing Line", 1)
    add_para(
        doc,
        "EventSlot can honestly be presented as the registration, attendance, and event-operations layer for Hallelujah Challenge. It is not accurate to sell it as a full live-broadcast platform or a fully automated multi-day volunteer system today, but it is accurate to say it can manage the core event workflow and grow from there.",
    )

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    footer = section2.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("EventSlot x Hallelujah Challenge | safe-claims version")
    set_run_font(run, size=9, color="666666")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
