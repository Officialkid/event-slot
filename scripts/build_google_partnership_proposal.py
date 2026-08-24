from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "EventSlot_Google_Partnership_Proposal_2026-08-20.docx"
SNAPSHOT = ROOT / "investor-package" / "data-room" / "traction_snapshot.json"

GOOGLE_SOURCES = [
    "Google Cloud Startup Perks: up to $200,000 in Google Cloud credits, up to $350,000 for AI startups, and other startup benefits.",
    "Cloud Run pricing: pay-per-use, always-free tier, 240,000 vCPU-seconds and 450,000 GiB-seconds monthly free tier, rounded to nearest 100 ms.",
    "Live Stream API pricing: HD H.264 input $0.14/hour, HD H.264 output $0.45/hour, distribution endpoint $0.75/hour, auto-captioning/translation $0.75/minute.",
    "Forms API: retrieve form contents and responses via Google Workspace Forms API, subject to authorization and EAP instructions.",
    "Meet API: create/manage meeting spaces, list participants, and retrieve recordings/transcripts/smart notes.",
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(
    doc: Document,
    text: str = "",
    *,
    size=11,
    color="000000",
    bold=False,
    italic=False,
    after=8,
    before=0,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    style=None,
):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color="2E74B5", bold=True)
    elif level == 2:
        set_run_font(run, size=13, color="2E74B5", bold=True)
    else:
        set_run_font(run, size=12, color="1F4D78", bold=True)


def add_bullets(doc: Document, items: list[str], style="List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color="000000")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color="000000")


def format_table(table, widths: list[float], header_fill="F2F4F7") -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.08
                for run in p.runs:
                    set_run_font(run, size=10.5, color="000000")
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def set_footer(section, left_text: str, right_text: str) -> None:
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = footer_p.add_run(left_text)
    set_run_font(run, size=9, color="666666")
    tab = footer_p.add_run("\t")
    set_run_font(tab, size=9, color="666666")
    run2 = footer_p.add_run(right_text)
    set_run_font(run2, size=9, color="666666")
    footer_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)


def add_cover(doc: Document, metrics: dict, round_terms: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("EVENTSLOT X GOOGLE")
    set_run_font(run, size=11, color="666666", bold=True)

    add_para(
        doc,
        "Strategic Partnership Proposal",
        size=26,
        color="000000",
        bold=False,
        after=4,
    )
    add_para(
        doc,
        "Cloud, Workspace, Maps, Meet, Forms, AI, and Live Broadcast for the event infrastructure layer",
        size=14,
        color="555555",
        after=14,
    )

    meta = doc.add_table(rows=0, cols=2)
    rows = [
        ("Prepared by", "EventSlot"),
        ("Prepared for", "Google Cloud, Google Workspace, and Google ecosystem partnerships"),
        ("Date", "August 20, 2026"),
        ("Partnership goal", "A meaningful infrastructure and ecosystem partnership, not a one-off credits request"),
    ]
    for key, value in rows:
        cells = meta.add_row().cells
        cells[0].text = key
        cells[1].text = value
    format_table(meta, [1.5, 5.0], header_fill="F7F9FC")
    for row in meta.rows:
        shade_cell(row.cells[0], "E8EEF5")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    cell = callout.rows[0].cells[0]
    set_cell_width(cell, 6.5)
    shade_cell(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(
        "EventSlot is already a working event operations platform. We are not asking Google to fund an idea from zero. We are proposing a partnership that turns Google infrastructure and APIs into a practical event experience for African organizers, starting with churches, conferences, and summits."
    )
    set_run_font(run, size=11.5, color="1F3A5F", bold=True)


def build_document() -> Path:
    with open(SNAPSHOT, "r", encoding="utf-8") as f:
        snapshot = json.load(f)

    metrics = snapshot["metrics"]
    round_terms = snapshot["round"]
    funds = snapshot.get("use_of_funds", [])

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    add_cover(doc, metrics, round_terms)
    doc.add_paragraph().add_run("")
    add_para(doc, " ", after=6)

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "EventSlot is building an event operations layer for organizers who need more than a form, more than a ticket page, and more control than a spreadsheet. The product already handles public event creation, capacity limits, waitlists, dynamic promotion, event-day verification, team access, analytics, exports, and organizer communications. The next strategic step is to deepen that platform with Google technology so EventSlot can support registrations, calendars, maps, meetings, AI-assisted operations, and hybrid live broadcast in one coherent workflow.",
    )
    add_para(
        doc,
        "This proposal is intentionally not framed as a request for charity or a one-off sponsorship. We are looking for a meaningful partnership: startup support where eligible, technical guidance, ecosystem access, and room to pilot an event infrastructure stack on Google Cloud and adjacent Google products. In return, Google gets a practical event-tech use case, cloud consumption, Workspace and Maps adoption, and a credible growth story in an African market that is still early but already active.",
    )

    add_heading(doc, "Verified Traction Snapshot", 1)
    traction = doc.add_table(rows=1, cols=4)
    traction.rows[0].cells[0].text = "Metric"
    traction.rows[0].cells[1].text = "Value"
    traction.rows[0].cells[2].text = "Metric"
    traction.rows[0].cells[3].text = "Value"
    traction_pairs = [
        ("Users", str(metrics["users"])),
        ("Events", str(metrics["events"])),
        ("Registrations", str(metrics["registrations"])),
        ("Active events", str(metrics["active_events"])),
        ("Organizers with events", str(metrics["organizers_with_events"])),
        ("Organizers with 30+ confirmed on an event", str(metrics["organizers_with_30_plus_confirmed"])),
        ("Organizers created in the last 30 days", str(metrics["organizers_last_30_days"])),
        ("Pro users / free users", f'{metrics["pro_users"]} / {metrics["free_users"]}'),
    ]
    for i in range(0, len(traction_pairs), 2):
        row = traction.add_row().cells
        left = traction_pairs[i]
        right = traction_pairs[i + 1] if i + 1 < len(traction_pairs) else ("", "")
        row[0].text, row[1].text = left
        row[2].text, row[3].text = right
    format_table(traction, [1.6, 1.7, 2.9, 0.9])

    add_para(
        doc,
        f"Latest verified repo snapshot in the workspace: {metrics['as_of']}. The top active events include {metrics['top_events'][0]['title']} ({metrics['top_events'][0]['confirmed_count']} confirmed), {metrics['top_events'][1]['title']} ({metrics['top_events'][1]['confirmed_count']} confirmed), and other live events with meaningful attendance pressure.",
    )
    add_para(
        doc,
        f"Current monetization is still early: report revenue is KSh {metrics['report_revenue_ksh']:,} and report downloads are {metrics['report_downloads']}. That matters because it tells the story honestly. EventSlot has usage and operational relevance already; the next step is to convert that into stronger, repeatable commercial value.",
    )

    add_heading(doc, "What EventSlot Already Does Well", 1)
    add_bullets(
        doc,
        [
            "Public event creation with custom registration questions and shareable links.",
            "Capacity management with confirmed-count tracking and remaining-spots visibility.",
            "Automatic waitlist handling when capacity is reached.",
            "FIFO promotion when capacity is increased.",
            "Organiser dashboard with activity, pressure, and event health views.",
            "Registration exports, duplicate detection, manual registration, and attendee management.",
            "Ticket verification, check-in, verifier-only access, and walk-in support.",
            "Team invites and per-event access assignment.",
            "Event analytics, insights, and feedback collection.",
            "PWA support for installable, mobile-friendly usage.",
            "Payments and ticket tiers for paid events.",
            "Consent, privacy, and KDPA-aligned controls.",
            "Google Calendar integration and Google Maps direction support.",
            "FAQ, WhatsApp/community link, and attendee contact pathways.",
        ],
    )

    add_heading(doc, "Where Google Fits", 1)
    scope = doc.add_table(rows=1, cols=4)
    scope.rows[0].cells[0].text = "Google capability"
    scope.rows[0].cells[1].text = "EventSlot use"
    scope.rows[0].cells[2].text = "Partner value"
    scope.rows[0].cells[3].text = "Status"
    scope_rows = [
        ("Cloud Run", "Serve the app and burst traffic on launch days", "Pay-per-use compute that matches event spikes", "Live architecture target"),
        ("Cloud Storage / Drive", "Store event assets, reports, and post-event files", "Storage and document workflows for organizers", "Proposed / partial"),
        ("Maps Platform", "Venue directions and location discovery", "Every physical event needs location utility", "Live / partial"),
        ("Calendar", "Schedule sync and reminders", "Better attendance continuity", "Live / partial"),
        ("Meet", "Virtual meetings, staff sessions, and post-event artifacts", "Meeting infrastructure for remote participation", "Proposed"),
        ("Forms API", "Import existing form registrations into EventSlot", "Migration path from basic collection to operations", "Proposed"),
        ("Live Stream API", "EventSlot Live for hybrid broadcast", "Video infrastructure without building a video stack", "Proposed"),
        ("Gemini / AI", "Summaries, FAQs, insights, and content drafting", "AI adoption in a practical event workflow", "Proposed"),
    ]
    for row_data in scope_rows:
        row = scope.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    format_table(scope, [1.2, 2.0, 2.2, 1.1])

    add_heading(doc, "The Live Broadcast Opportunity", 1)
    add_para(
        doc,
        "EventSlot Live should be treated as a hybrid-event layer, not as a replacement for a streaming platform from day one. The idea is simple: EventSlot owns registration, ticketing, access, and the attendee journey. Google handles the heavy broadcast infrastructure through Live Stream API, while Google Meet can support internal organizer sessions and post-event artifacts where appropriate.",
    )
    add_para(
        doc,
        "That distinction matters. We are not proposing that EventSlot build YouTube or Zoom from scratch. We are proposing a controlled, usage-based broadcast path that lets a church conference, summit, or leadership gathering serve both a physical room and a remote audience in one product surface.",
    )

    broadcast = doc.add_table(rows=1, cols=4)
    broadcast.rows[0].cells[0].text = "Scenario"
    broadcast.rows[0].cells[1].text = "Google components"
    broadcast.rows[0].cells[2].text = "Illustrative cost"
    broadcast.rows[0].cells[3].text = "Why it matters"
    broadcast_rows = [
        ("Pilot live broadcast", "1 HD input + 1 HD output", "$0.59/hour", "Base video cost before any distribution or captions"),
        ("Pilot with distribution endpoint", "1 HD input + 1 HD output + 1 distribution endpoint", "$1.34/hour", "More realistic for remote delivery architecture"),
        ("Two-hour broadcast with distribution", "Same as above, for 2 hours", "$2.68 total", "A simple, credible broadcast example for a church summit"),
        ("Auto-captions or translation", "$0.75/minute", "$45/hour", "Should be a premium add-on, not the default"),
    ]
    for row_data in broadcast_rows:
        row = broadcast.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    format_table(broadcast, [1.35, 2.1, 1.2, 1.85])
    add_para(
        doc,
        "Google Cloud also documents a pay-per-use Cloud Run model with an always-free tier, which is exactly the kind of economics that fit EventSlot's traffic pattern: quiet most of the time, then suddenly busy when registrations open or an event goes live.",
    )

    add_heading(doc, "Partnership Economics", 1)
    add_para(
        doc,
        "A meaningful partnership should not be only about money, and it should not be only about software access. The value exchange should work in both directions. Google helps EventSlot scale with infrastructure, startup credits, and technical support. EventSlot helps Google with cloud consumption, Workspace adoption, Maps usage, hybrid-event use cases, and a credible growth story in an African market.",
    )

    benefit = doc.add_table(rows=1, cols=3)
    benefit.rows[0].cells[0].text = "What Google can offer"
    benefit.rows[0].cells[1].text = "Why it helps EventSlot"
    benefit.rows[0].cells[2].text = "What Google gets"
    benefit_rows = [
        ("Startup credits", "Lower infrastructure burn while we harden the platform", "A startup customer that can scale with Google Cloud"),
        ("Workspace / Maps / AI access", "Faster product integration and better organizer workflows", "Real-world adoption of the Google ecosystem"),
        ("Forms, Meet, and Live Stream support", "A practical path from registration to virtual and hybrid events", "A sticky event-tech use case"),
        ("Technical support and architecture guidance", "Less trial and error on critical paths", "A better chance of a durable reference customer"),
    ]
    for row_data in benefit_rows:
        row = benefit.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
    format_table(benefit, [1.8, 2.55, 2.15])

    add_heading(doc, "Current Google References For This Proposal", 1)
    add_numbers(
        doc,
        [
            "Google Cloud Startup Perks and startup program benefits, including cloud credits and Maps/Workspace support where eligible.",
            "Cloud Run pricing and pay-per-use model, including the always-free tier.",
            "Live Stream API pricing for HD input/output and distribution endpoints.",
            "Google Workspace Forms API for retrieving form contents and responses.",
            "Google Meet REST API for meeting spaces, participants, and artifacts.",
        ],
    )

    add_heading(doc, "Proposed Partnership Ask", 1)
    add_bullets(
        doc,
        [
            "Google Cloud startup credits or equivalent startup support where EventSlot qualifies.",
            "Technical architecture support for Cloud Run, storage, security, and scaling.",
            "Maps Platform support for venue and directions usage.",
            "Workspace integration guidance for Calendar, Drive, and Forms.",
            "Meet and Live Stream API guidance for virtual and hybrid event workflows.",
            "AI collaboration for event summaries, FAQs, support automation, and post-event insights.",
            "A pilot-friendly relationship that can grow into a stronger ecosystem partnership as traction expands.",
        ],
    )

    add_heading(doc, "What EventSlot Can Commit To", 1)
    add_bullets(
        doc,
        [
            "Build on Google Cloud where it makes technical and commercial sense.",
            "Use Google products as part of the event workflow, not as decorative integrations.",
            "Track and report cloud consumption, Maps usage, Workspace usage, and hybrid-event metrics.",
            "Maintain a provider-neutral architecture where practical so the business stays resilient.",
            "Treat the partnership as a long-term relationship, not a one-off discount conversation.",
        ],
    )

    add_heading(doc, "Suggested Next Steps", 1)
    add_numbers(
        doc,
        [
            "Approve the partnership narrative and the exact Google ask.",
            "Validate the live broadcast pilot architecture on Cloud Run + Live Stream API.",
            "Map the Forms import path for registration migration use cases.",
            "Prepare a short one-pager and an architecture brief for Google teams.",
            "Package the traction snapshot and cost model into a partner-ready leave-behind.",
        ],
    )

    add_heading(doc, "Appendix: Product Snapshot", 1)
    add_para(
        doc,
        "The proposal rests on the current repository implementation. EventSlot already has a live event creation flow, capacity and waitlist management, dynamic promotion, organiser dashboards, analytics, notifications, exports, ticket verification, QR/check-in workflows, team access, Google Calendar support, Maps direction support, FAQ surfaces, consent controls, and PWA behavior. The live broadcast layer is the next proposed extension, not a claim that the feature is already fully built.",
    )
    add_para(
        doc,
        "This is why the partnership is meaningful: Google would not just be backing an abstract idea. Google would be supporting a working event platform that is already in motion and that can grow into a broader event-infrastructure layer with the right support.",
    )

    add_heading(doc, "Appendix: Official Google Sources Consulted", 1)
    add_bullets(doc, GOOGLE_SOURCES)

    if funds:
        add_heading(doc, "Appendix: Current Use of Funds Snapshot", 1)
        uf = doc.add_table(rows=1, cols=3)
        uf.rows[0].cells[0].text = "Area"
        uf.rows[0].cells[1].text = "Amount"
        uf.rows[0].cells[2].text = "Why it matters"
        for row_data in funds:
            row = uf.add_row().cells
            row[0].text, row[1].text, row[2].text = row_data
        format_table(uf, [1.8, 1.1, 3.6])

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_footer(section, "EventSlot | Google partnership proposal", "1")
    set_footer(section2, "EventSlot | Google partnership proposal", "2")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
