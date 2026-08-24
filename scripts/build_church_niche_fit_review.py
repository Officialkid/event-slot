from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "EventSlot_Church_Summit_Niche_Fit_Review_2026-08-06.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None, size=11, bold=False,
                  color="000000", after=8, before=0, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color="000000")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(item)
        set_run_font(run, size=11, color="000000")


def format_table(table, widths: list[float], header_fill="F2F4F7"):
    table.style = "Table Grid"
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, color="000000")
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    size = 16 if level == 1 else 13
    color = "2E74B5" if level <= 2 else "1F4D78"
    set_run_font(run, size=size, color=color, bold=True)


def add_metadata(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    format_table(table, [1.6, 4.9], header_fill="FFFFFF")
    for row in table.rows:
        set_cell_shading(row.cells[0], "E8EEF5")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    add_paragraph(doc, "EventSlot Niche Fit Review", size=24, bold=False, after=2)
    add_paragraph(
        doc,
        "Church Conferences, Christian Summits, and Faith-Based Event Operations",
        size=14,
        color="555555",
        after=16,
    )
    add_metadata(
        doc,
        [
            ("Prepared for", "EventSlot product and strategy team"),
            ("Date", "August 6, 2026"),
            ("Primary lens", "Christian event organiser running conferences, summits, services, and volunteer-led gatherings"),
            ("Bottom line", "EventSlot is already strong for attendee registration and event-day control, but still underbuilt for volunteer management workflows."),
        ],
    )

    add_heading(doc, "Executive Summary", 1)
    add_paragraph(
        doc,
        "EventSlot already covers many of the operational jobs a church or summit organiser cares about after the event is announced: custom registration forms, capacity tracking, waitlists, public event pages, team access, event-day verification, exports, analytics, and feedback. The product is therefore not misaligned with the Christian events niche. In fact, it already has useful church-facing signals such as a faith-event template, WhatsApp and community links, consent controls, maps, FAQs, and verifier access for gate teams.",
    )
    add_paragraph(
        doc,
        "The deeper issue is that the current platform is attendee-centric rather than volunteer-centric. A church conference usually starts with volunteer recruitment, approval, role assignment, communication, and on-ground coordination before normal attendee registration becomes the main operating surface. EventSlot does not yet model volunteers as a first-class workflow. It can collect volunteer data through custom questions, but it does not yet support role quotas, approvals, service-team assignments, volunteer-specific communication, or volunteer check-in as dedicated product features.",
    )
    add_paragraph(
        doc,
        "That means the product is good enough to serve churches and summits today, especially where the organiser mainly needs registration and attendance control. It is not yet ideal for churches that run ministry teams, ushers, protocol teams, media volunteers, security teams, prayer teams, and multi-stage event operations from one central workspace.",
    )

    add_heading(doc, "Evaluation Lens", 1)
    add_numbered(
        doc,
        [
            "A Christian organiser often recruits volunteers before opening general attendee registration.",
            "Volunteer roles usually have limited slots, different responsibilities, and an approval process rather than instant confirmation.",
            "Church conferences and summits usually rely heavily on WhatsApp, community links, service teams, event-day verification, and post-event follow-up.",
            "The strongest niche-fit product is not only a registration form. It is an operating system for pre-event coordination, event-day flow, and follow-up.",
        ],
    )

    add_heading(doc, "Current Fit Scorecard", 1)
    score_table = doc.add_table(rows=1, cols=4)
    score_table.rows[0].cells[0].text = "Capability"
    score_table.rows[0].cells[1].text = "Score / 5"
    score_table.rows[0].cells[2].text = "Assessment"
    score_table.rows[0].cells[3].text = "Why it matters for churches and summits"
    score_rows = [
        ("Attendee registration", "4.5", "Strong", "Core forms, question builder, public event page, and attendee self-service are already in place."),
        ("Capacity and overflow handling", "4.5", "Strong", "Important for limited-seat conferences, breakout sessions, and high-demand summits."),
        ("Event-day check-in and verification", "4.0", "Strong", "Verifier-only ticket workspace and check-in routes support controlled entry."),
        ("Organizer collaboration", "3.5", "Good", "Team invites and per-event access exist, but permissions are not yet deeply role-based."),
        ("Communications", "3.0", "Moderate", "Email campaigns work, but WhatsApp and SMS are still not first-class delivery channels."),
        ("Church community continuity", "3.5", "Good", "Community links and WhatsApp contact help, but follow-up journeys are still light."),
        ("Volunteer recruitment and operations", "1.5", "Weak", "No first-class volunteer objects, role slots, approvals, or volunteer dashboards."),
        ("Conference programme complexity", "2.5", "Developing", "General event support is good, but session, speaker, and track operations are still missing."),
    ]
    for row in score_rows:
        cells = score_table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(score_table, [1.7, 0.8, 1.2, 2.8])

    add_heading(doc, "What EventSlot Already Does Well For This Niche", 1)
    add_heading(doc, "1. It already supports conference and church-style event creation", 2)
    add_paragraph(
        doc,
        "The product includes event templates for both \"Conference\" and \"Church / Faith Event.\" That is a meaningful signal because it shows the system already recognizes that event formats are not identical. The church template also asks first-time-visitor information, which is especially relevant for church gatherings and ministry follow-up.",
    )
    add_heading(doc, "2. It is strong on attendee registration operations", 2)
    add_bullets(
        doc,
        [
            "Custom registration questions can be added during event setup.",
            "Attendees do not need EventSlot accounts for normal registration.",
            "Duplicate detection, registration editing, and manual registration routes are present.",
            "Confirmed and waitlisted registrants can be exported in CSV, PDF, and Excel-friendly formats.",
        ],
    )
    add_heading(doc, "3. Capacity and waitlist handling are a real advantage", 2)
    add_paragraph(
        doc,
        "For Christian conferences and summits, space pressure is common. EventSlot already tracks confirmed count, waitlist count, and capacity, and it promotes people from the waitlist when capacity increases. That is stronger than a simple form workflow and especially useful for ministry events where late seat releases are normal.",
    )
    add_heading(doc, "4. Event-day control is already credible", 2)
    add_bullets(
        doc,
        [
            "Verifier access exists for gate teams and temporary event workers.",
            "Ticket verification and entry logging are already modelled in the system.",
            "Walk-in event support exists for on-site check-in workflows.",
            "Checked-in status is reflected in registration and analytics surfaces.",
        ],
    )
    add_heading(doc, "5. Church-friendly distribution and follow-up signals already exist", 2)
    add_bullets(
        doc,
        [
            "Community links can point people into WhatsApp or Telegram groups.",
            "WhatsApp contact mode can open a prefilled organiser chat.",
            "Maps and directions links help physical gatherings.",
            "FAQs, remaining-spots display, and attendee consent settings improve public event clarity.",
        ],
    )
    add_heading(doc, "6. Organizer teamwork is present", 2)
    add_paragraph(
        doc,
        "The team workspace allows invites, accepted members, and event-by-event access assignment. For a church or summit environment, this is valuable for separating central organisers from event helpers, though the permission model is still not deep enough for ministry-specific operations.",
    )
    add_heading(doc, "7. Reporting and insight surfaces are ahead of many early-stage tools", 2)
    add_bullets(
        doc,
        [
            "Event analytics track views, registrations, conversion rate, check-in rate, waitlist movement, and source breakdown.",
            "Feedback collection exists for post-event learning.",
            "AI insights and exports already support after-action review.",
        ],
    )

    add_heading(doc, "Where EventSlot Is Weak For Church and Summit Organisers", 1)
    weakness_items = [
        "Volunteers are not a first-class object in the system. Today they are only approximated through generic custom questions or demo content.",
        "There is no separate volunteer funnel before attendee registration. Churches often recruit service teams first, then attendees second.",
        "There are no volunteer role capacities such as ushers: 20, protocol: 8, intercession: 12, media: 6, security support: 10.",
        "There is no volunteer approval flow. A ministry lead cannot review applicants and accept some while waitlisting or declining others.",
        "There is no service-team assignment layer. The system cannot say who is on welcome, altar ministry, parking, registration desk, worship support, or technical operations.",
        "There is no volunteer-only communication layer for reminders, reporting times, dress code, or supervisor instructions.",
        "Conference programme management is still light. There is no session schedule, breakout management, speaker roster, or track-level capacity model.",
        "Communications remain email-first. The docs explicitly say SMS and WhatsApp delivery are roadmap items, which matters because many church workflows are WhatsApp-heavy.",
        "Team access exists, but fine-grained permissions are still limited. For example, one person may need check-in rights but not billing or event editing rights.",
        "There is no church-specific post-event follow-up workflow such as first-time visitor follow-up, prayer requests, ministry interest capture, or campus/cell-group routing.",
    ]
    add_bullets(doc, weakness_items)

    add_heading(doc, "What The Product Currently Favors", 1)
    comparison = doc.add_table(rows=1, cols=3)
    comparison.rows[0].cells[0].text = "Organizer type"
    comparison.rows[0].cells[1].text = "How well EventSlot fits today"
    comparison.rows[0].cells[2].text = "Reason"
    comparison_rows = [
        ("Generic event organizer", "High", "The current system is excellent for creating an event, collecting registrations, handling capacity, checking people in, and exporting lists."),
        ("Church conference organizer", "Medium-high", "The core attendee workflow is good, but volunteer and ministry-team coordination still need dedicated product depth."),
        ("Summit organizer with multiple teams", "Medium", "The platform handles registrations and entry well, but complex programme, speaker, and service-team workflows are still underbuilt."),
        ("Volunteer-heavy ministry organizer", "Low-medium", "Possible through workarounds, but not yet supported as a clean first-class operational flow."),
    ]
    for row in comparison_rows:
        cells = comparison.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(comparison, [1.8, 1.1, 3.6])

    add_heading(doc, "Highest-Priority Features To Implement Next", 1)
    add_heading(doc, "Priority 1: Volunteer management foundation", 2)
    add_numbered(
        doc,
        [
            "Add a volunteer mode for events, separate from attendee registration.",
            "Add volunteer role slots and capacities by function or ministry area.",
            "Add volunteer approval states such as pending, approved, waitlisted, declined, and checked-in.",
            "Add volunteer-specific exports and dashboards.",
            "Add volunteer reminder templates for reporting time, dress code, arrival gate, and team lead contact.",
        ],
    )
    add_heading(doc, "Priority 2: Church and summit operations depth", 2)
    add_numbered(
        doc,
        [
            "Add session, breakout, or service-segment management for conferences and summits.",
            "Add speaker, minister, or guest-host records with schedule attachment.",
            "Add more granular staff permissions such as check-in only, communications only, reporting only, and full organiser access.",
            "Add bulk attendee segmentation such as first-time visitors, leaders, volunteers, speakers, and general attendees.",
            "Add stronger post-event ministry follow-up flows for first-time guests and decision follow-up.",
        ],
    )
    add_heading(doc, "Priority 3: Communication channels that match the niche", 2)
    add_numbered(
        doc,
        [
            "Ship first-class WhatsApp outbound workflows where policy and integrations allow.",
            "Add SMS fallback for critical event reminders.",
            "Create message templates tuned for churches and summits, not only generic event language.",
        ],
    )

    add_heading(doc, "Suggested Product Positioning For Now", 1)
    add_paragraph(
        doc,
        "The strongest honest positioning today is this: EventSlot is already a strong operating platform for church conferences, summits, and faith-based gatherings that need registration, capacity control, event-day verification, and organiser coordination. It is not yet the full operating system for volunteer-led ministry events, but it is close enough that a focused volunteer module could make the niche much more defensible.",
    )
    add_paragraph(
        doc,
        "In other words, the system can already serve the niche, but the niche story becomes much stronger once EventSlot supports the real pre-event volunteer journey instead of only the public attendee journey.",
    )

    add_heading(doc, "Recommended 90-Day Product Direction", 1)
    roadmap = doc.add_table(rows=1, cols=4)
    roadmap.rows[0].cells[0].text = "Window"
    roadmap.rows[0].cells[1].text = "Focus"
    roadmap.rows[0].cells[2].text = "Key deliverables"
    roadmap.rows[0].cells[3].text = "Expected outcome"
    roadmap_rows = [
        ("0-30 days", "Volunteer foundation", "Volunteer application form, role slots, approval statuses, role-based exports", "Church organisers can recruit and classify teams inside EventSlot instead of outside it."),
        ("31-60 days", "Operations depth", "Volunteer reminders, role assignments, team lead view, volunteer check-in", "The product begins to support ministry execution, not only public registration."),
        ("61-90 days", "Niche polish", "Church/summit templates, segmented follow-up, WhatsApp-first touchpoints, conference schedule layer", "EventSlot starts to feel intentionally built for the chosen niche."),
    ]
    for row in roadmap_rows:
        cells = roadmap.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(roadmap, [0.9, 1.1, 2.2, 2.3])

    add_heading(doc, "Final Verdict", 1)
    add_paragraph(
        doc,
        "EventSlot is already more useful for church conferences and Christian summits than it might first appear. The current feature set is not a mismatch. It already handles many of the things these organisers need once the event is open to the public. The main gap is not event management in general; it is volunteer operations specifically.",
    )
    add_paragraph(
        doc,
        "So if the goal is to focus on Christian conferences, summits, and church events as a starting niche, the decision makes sense. The next smart move is to build the volunteer layer quickly and deliberately. Once that exists, EventSlot will move from being a strong registration platform for church events to being a much stronger church event operations platform.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix: Repository Evidence Reviewed", 1)
    add_bullets(
        doc,
        [
            "Event templates include both conference and church / faith-event presets in lib/eventTemplates.ts.",
            "The event model supports capacity, waitlist, community links, WhatsApp contact, maps, FAQs, event type, team access, email campaigns, feedback, and analytics in prisma/schema.prisma.",
            "The organiser create and event dashboard flows include join opens time, map directions, attendee consent, community link, WhatsApp contact mode, and remaining-spots controls.",
            "Team workspace pages and APIs support member invites and per-event access assignment.",
            "Verifier access, verify-ticket routes, walk-in check-in routes, and entry logs support event-day control.",
            "Email campaigns, exports, analytics, duplicates, feedback, and AI insights are already present in organiser surfaces and APIs.",
            "The communications documentation states that EventSlot is currently email-first and that WhatsApp and SMS are future integrations.",
            "Only light volunteer wording appears in verifier copy and demo data, which reinforces that volunteer operations are not yet a first-class product area.",
        ],
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(path)
